import asyncio
import csv
import io
import os
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

UNSUPPORTED_BINARY_EXTENSIONS = {
	'png',
	'jpg',
	'jpeg',
	'gif',
	'bmp',
	'svg',
	'webp',
	'ico',
	'mp3',
	'mp4',
	'wav',
	'avi',
	'mov',
	'zip',
	'tar',
	'gz',
	'rar',
	'exe',
	'bin',
	'dll',
	'so',
}


def _build_filename_error_message(file_name: str, supported_extensions: list[str]) -> str:
	"""Build a specific error message explaining why the filename was rejected and how to fix it."""
	base = os.path.basename(file_name)

	# Check for binary/image extension
	if '.' in base:
		_, ext = base.rsplit('.', 1)
		ext_lower = ext.lower()
		if ext_lower in UNSUPPORTED_BINARY_EXTENSIONS:
			return (
				f"Error: Cannot write binary/image file '{base}'. "
				f'The write_file tool only supports text-based files. '
				f'Supported extensions: {", ".join("." + e for e in supported_extensions)}. '
				f'For screenshots, the browser automatically captures them - do not try to save screenshots as files.'
			)
		if ext_lower not in supported_extensions:
			return (
				f"Error: Unsupported file extension '.{ext_lower}' in '{base}'. "
				f'Supported extensions: {", ".join("." + e for e in supported_extensions)}. '
				f'Please rename the file to use a supported extension.'
			)

	# No extension or no dot
	if '.' not in base:
		return (
			f"Error: Filename '{base}' has no extension. "
			f'Please add a supported extension: {", ".join("." + e for e in supported_extensions)}.'
		)

	return (
		f"Error: Invalid filename '{base}'. "
		f'Filenames must contain only letters, numbers, underscores, hyphens, dots, parentheses, and spaces. '
		f'Supported extensions: {", ".join("." + e for e in supported_extensions)}.'
	)


DEFAULT_FILE_SYSTEM_PATH = 'browseruse_agent_data'


class FileSystemError(Exception):
	"""Custom exception for file system operations that should be shown to LLM"""

	pass


class BaseFile(BaseModel, ABC):
	"""Base class for all file types"""

	name: str
	content: str = ''

	# --- Subclass must define this ---
	@property
	@abstractmethod
	def extension(self) -> str:
		"""File extension (e.g. 'txt', 'md')"""
		pass

	def write_file_content(self, content: str) -> None:
		"""Update internal content (formatted)"""
		self.update_content(content)

	def append_file_content(self, content: str) -> None:
		"""Append content to internal content"""
		self.update_content(self.content + content)

	# --- These are shared and implemented here ---

	def update_content(self, content: str) -> None:
		self.content = content

	def sync_to_disk_sync(self, path: Path) -> None:
		file_path = path / self.full_name
		file_path.write_text(self.content, encoding='utf-8')

	async def sync_to_disk(self, path: Path) -> None:
		file_path = path / self.full_name
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: file_path.write_text(self.content, encoding='utf-8'))

	async def write(self, content: str, path: Path) -> None:
		self.write_file_content(content)
		await self.sync_to_disk(path)

	async def append(self, content: str, path: Path) -> None:
		self.append_file_content(content)
		await self.sync_to_disk(path)

	def read(self) -> str:
		return self.content

	@property
	def full_name(self) -> str:
		return f'{self.name}.{self.extension}'

	@property
	def get_size(self) -> int:
		return len(self.content)

	@property
	def get_line_count(self) -> int:
		return len(self.content.splitlines())


class MarkdownFile(BaseFile):
	"""Markdown file implementation"""

	@property
	def extension(self) -> str:
		return 'md'


class TxtFile(BaseFile):
	"""Plain text file implementation"""

	@property
	def extension(self) -> str:
		return 'txt'


class JsonFile(BaseFile):
	"""JSON file implementation"""

	@property
	def extension(self) -> str:
		return 'json'


class CsvFile(BaseFile):
	"""CSV file implementation with automatic RFC 4180 normalization.

	LLMs frequently produce malformed CSV (missing quotes around fields with commas,
	inconsistent empty fields, unescaped internal quotes). This class parses the raw
	content through Python's csv module on every write to guarantee well-formed output.
	"""

	@property
	def extension(self) -> str:
		return 'csv'

	@staticmethod
	def _normalize_csv(raw: str) -> str:
		"""Parse and re-serialize CSV content to fix quoting, empty fields, and escaping.

		Handles common LLM mistakes: unquoted fields containing commas,
		unescaped quotes inside fields, inconsistent empty fields,
		trailing/leading blank lines, and double-escaped JSON output
		(literal backslash-n and backslash-quote instead of real newlines/quotes).
		"""
		stripped = raw.strip('\n\r')
		if not stripped:
			return raw

		# Detect double-escaped LLM tool call output: if the content has no real
		# newlines but contains literal \n sequences, the entire string is likely
		# double-escaped JSON. Unescape \" → " first, then \n → newline.
		if '\n' not in stripped and '\\n' in stripped:
			stripped = stripped.replace('\\"', '"')
			stripped = stripped.replace('\\n', '\n')

		reader = csv.reader(io.StringIO(stripped))
		rows: list[list[str]] = []
		for row in reader:
			# Skip completely empty rows (artifacts of blank lines)
			if row:
				rows.append(row)

		if not rows:
			return raw

		out = io.StringIO()
		writer = csv.writer(out, lineterminator='\n')
		writer.writerows(rows)
		# Strip trailing newline so callers (write_file action) control line endings
		return out.getvalue().rstrip('\n')

	def write_file_content(self, content: str) -> None:
		"""Normalize CSV content before storing."""
		self.update_content(self._normalize_csv(content))

	def append_file_content(self, content: str) -> None:
		"""Normalize the appended CSV rows and merge with existing content."""
		normalized_new = self._normalize_csv(content)
		if not normalized_new.strip('\n\r'):
			return
		existing = self.content
		if existing and not existing.endswith('\n'):
			existing += '\n'
		combined = existing + normalized_new
		self.update_content(self._normalize_csv(combined))


class JsonlFile(BaseFile):
	"""JSONL (JSON Lines) file implementation"""

	@property
	def extension(self) -> str:
		return 'jsonl'


class PdfFile(BaseFile):
	"""PDF file implementation"""

	@property
	def extension(self) -> str:
		return 'pdf'

	def sync_to_disk_sync(self, path: Path) -> None:
		# Lazy import reportlab
		from reportlab.lib.pagesizes import letter
		from reportlab.lib.styles import getSampleStyleSheet
		from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

		file_path = path / self.full_name
		try:
			# Create PDF document
			doc = SimpleDocTemplate(str(file_path), pagesize=letter)
			styles = getSampleStyleSheet()
			story = []

			# Convert markdown content to simple text and add to PDF
			# For basic implementation, we'll treat content as plain text
			# This avoids the AGPL license issue while maintaining functionality
			content_lines = self.content.split('\n')

			for line in content_lines:
				if line.strip():
					# Handle basic markdown headers
					if line.startswith('# '):
						para = Paragraph(line[2:], styles['Title'])
					elif line.startswith('## '):
						para = Paragraph(line[3:], styles['Heading1'])
					elif line.startswith('### '):
						para = Paragraph(line[4:], styles['Heading2'])
					else:
						para = Paragraph(line, styles['Normal'])
					story.append(para)
				else:
					story.append(Spacer(1, 6))

			doc.build(story)
		except Exception as e:
			raise FileSystemError(f"Error: Could not write to file '{self.full_name}'. {str(e)}")

	async def sync_to_disk(self, path: Path) -> None:
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: self.sync_to_disk_sync(path))


class DocxFile(BaseFile):
	"""DOCX file implementation"""

	@property
	def extension(self) -> str:
		return 'docx'

	def sync_to_disk_sync(self, path: Path) -> None:
		file_path = path / self.full_name
		try:
			from docx import Document

			doc = Document()

			# Convert content to DOCX paragraphs
			content_lines = self.content.split('\n')

			for line in content_lines:
				if line.strip():
					# Handle basic markdown headers
					if line.startswith('# '):
						doc.add_heading(line[2:], level=1)
					elif line.startswith('## '):
						doc.add_heading(line[3:], level=2)
					elif line.startswith('### '):
						doc.add_heading(line[4:], level=3)
					else:
						doc.add_paragraph(line)
				else:
					doc.add_paragraph()  # Empty paragraph for spacing

			doc.save(str(file_path))
		except Exception as e:
			raise FileSystemError(f"Error: Could not write to file '{self.full_name}'. {str(e)}")

	async def sync_to_disk(self, path: Path) -> None:
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: self.sync_to_disk_sync(path))


class HtmlFile(BaseFile):
	"""HTML file implementation"""

	@property
	def extension(self) -> str:
		return 'html'


class XmlFile(BaseFile):
	"""XML file implementation"""

	@property
	def extension(self) -> str:
		return 'xml'


class FileSystemState(BaseModel):
	"""Serializable state of the file system"""

	files: dict[str, dict[str, Any]] = Field(default_factory=dict)  # full filename -> file data
	base_dir: str
	extracted_content_count: int = 0
